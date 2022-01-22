import os
#import magic
import urllib.request
from flask import Flask, flash, request,send_file, redirect, render_template
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt


ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'png', 'jpg', 'jpeg', 'docx'])
UPLOAD_FOLDER = 'home/user/uploads'

def allowed_file(filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
app= Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = "secret key"
@app.route('/')
def index():
        return render_template("assignment.html")
@app.route('/hello')
def upload_form():
        return render_template('upload.html')

@app.route('/uploader', methods=['POST'])
def upload_file():
        if request.method == 'POST':
        # check if the post request has the file part
                if 'file' not in request.files:
                        flash('No file part')
                        return redirect(request.url)
                file = request.files['file']
                if file.filename == '':
                        flash('No file selected for uploading')
                        return redirect(request.url)
                if file and allowed_file(file.filename):
                        filename = secure_filename(file.filename)
                        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                        flash('File successfully uploaded')
                        return redirect('/hello')
                else:
                        flash('Allowed file types are txt, pdf, png, jpg, jpeg, gif')
                        return redirect(request.url)

@app.route('/assignment',methods=['POST'])
def form_example():
    # format user[name,roll no,class,CMPN]
    name = request.form.get('name')
    roll = request.form.get('roll')
    class1 = request.form.get('class')
    batch = request.form.get('batch')
    type = request.form.get('type')

    user = [name, roll, class1, batch]
    data = " Name:" + user[0] + "  ROLL-NO:" + user[1] + "  CLass:" + user[2] + "  Batch:" + user[3]
    data1 = str(data)
    print(data1)
    document = Document("sample.docx")
    # here int=2 defines level
    run = document.paragraphs[2].add_run()
    font = run.font
    # changing font
    font.name = 'Calibri'
    # to change size of text
    font.size = Pt(20)
    # making text bold
    font.bold = True
    # data to be passed
    run.text = data1
    document.save("result.docx")
    path = "result.docx"
    return send_file(path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
